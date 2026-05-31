#!/usr/bin/env python3
"""
wordModifier.py — In-place Word document modifier using python-docx.
Usage:
  python3.11 wordModifier.py <mode> <input.docx> <output.docx> <changes_json>
Modes:
  annotate_original  — highlight OLD values in YELLOW in the original document
  modify_green       — replace OLD values with NEW values, highlight NEW values in GREEN
  modify_clean       — replace OLD values with NEW values, NO highlights (clean download)
changes_json format:
  [{"fieldName": "Weight", "oldValue": "155", "newValue": "170", "unit": "g"}, ...]
Exit codes:
  0 = success, changes applied
  1 = error
  2 = no matches found (non-fatal)

Formatting preservation guarantee
----------------------------------
ALL original run formatting (font name, size, bold, italic, underline, color,
highlight, strike, etc.) is preserved exactly.  The ONLY things that change are:
  * The text content of the matched run (replaced with new value)
  * For annotate_original: a yellow <w:highlight> is added to the matched run
  * For modify_green:      a green  <w:highlight> is added to the matched run
  * For modify_clean:      no formatting change at all — pure text swap

The before/after text segments also keep their original per-run formatting
because we rebuild each segment run-by-run from the original run list.
"""
import sys
import json
import copy
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Highlight color constants (W3C names used by <w:highlight>)
YELLOW_HIGHLIGHT = "yellow"
GREEN_HIGHLIGHT  = "green"


def get_paragraph_full_text(para):
    """Get the full text of a paragraph across all runs."""
    return "".join(run.text for run in para.runs)


def _set_highlight_on_rPr(rPr, color_name):
    """Add/replace a <w:highlight> element in an rPr element."""
    for existing in rPr.findall(qn("w:highlight")):
        rPr.remove(existing)
    if color_name:
        h_el = OxmlElement("w:highlight")
        h_el.set(qn("w:val"), color_name)
        rPr.append(h_el)


def _apply_highlight_in_paragraph(para, search_term, highlight_color, replacement_text, keep_text):
    """
    Find search_term in the paragraph's runs and apply highlight/replacement
    while preserving ALL original run formatting.

    Strategy
    --------
    1. Build a character map that records which run each character belongs to.
    2. Find the match position in the concatenated text.
    3. Split the original run list into three logical segments:
         before  - characters before the match
         matched - characters of the match
         after   - characters after the match
    4. Rebuild the paragraph XML run-by-run:
         before/after: deep-copy each original run element, just update its text.
         matched:      deep-copy the original run element, update its text to the
                       replacement value, and (if needed) add/replace the highlight.
                       No other formatting property is touched.
    """
    runs = para.runs
    if not runs:
        return

    # 1. Build character map
    full_text = ""
    char_map = []  # (run_index, char_index_within_run)
    for ri, run in enumerate(runs):
        for ci, ch in enumerate(run.text):
            char_map.append((ri, ci))
            full_text += ch

    # 2. Find match
    pattern = re.compile(re.escape(search_term), re.IGNORECASE)
    match = pattern.search(full_text)
    if not match:
        return

    start, end = match.start(), match.end()

    # 3. Collect per-run slices
    # Walk through each position in the flat string, group by run, then split
    # each group by the before/matched/after segment boundaries.
    segments = []  # list of [run_idx, slice_text, segment_name]
    i = 0
    while i < len(full_text):
        run_idx = char_map[i][0]
        # Find end of this run's characters
        j = i
        while j < len(full_text) and char_map[j][0] == run_idx:
            j += 1
        # chars [i, j) all belong to run_idx; split by segment boundaries
        for seg_s, seg_e, seg_name in [
            (i,              min(j, start), "before"),
            (max(i, start),  min(j, end),   "matched"),
            (max(i, end),    j,              "after"),
        ]:
            if seg_s < seg_e:
                segments.append([run_idx, full_text[seg_s:seg_e], seg_name])
        i = j

    # Merge consecutive slices that share the same run index AND segment name
    merged = []
    for entry in segments:
        if merged and merged[-1][0] == entry[0] and merged[-1][2] == entry[2]:
            merged[-1][1] += entry[1]
        else:
            merged.append(list(entry))

    # 4. Rebuild paragraph XML
    p_elem = para._p
    original_run_elems = list(p_elem.findall(qn("w:r")))

    # Remove all existing <w:r> elements
    for r_elem in original_run_elems:
        p_elem.remove(r_elem)

    # Find insertion point (after <w:pPr> if present)
    insert_pos = 0
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        insert_pos = list(p_elem).index(pPr) + 1

    for run_idx, slice_text, seg_name in merged:
        if not slice_text:
            continue

        # Deep-copy the original run element — this preserves ALL formatting
        orig_r = original_run_elems[run_idx]
        new_r = copy.deepcopy(orig_r)

        # Replace the text content
        for t_el in new_r.findall(qn("w:t")):
            new_r.remove(t_el)
        t_el = OxmlElement("w:t")
        if seg_name == "matched" and not keep_text:
            t_el.text = replacement_text if replacement_text else slice_text
        else:
            t_el.text = slice_text
        if t_el.text.startswith(" ") or t_el.text.endswith(" "):
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(t_el)

        # For the matched segment, apply highlight ONLY — nothing else changes
        if seg_name == "matched" and highlight_color:
            rPr = new_r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                new_r.insert(0, rPr)
            _set_highlight_on_rPr(rPr, highlight_color)

        p_elem.insert(insert_pos, new_r)
        insert_pos += 1


def process_paragraph(para, changes, mode):
    """
    Process a single paragraph, applying changes based on mode.
    Returns True if any change was applied.
    """
    full_text = get_paragraph_full_text(para)
    if not full_text.strip():
        return False

    changed = False

    for change in changes:
        old_val = change.get("oldValue", "").strip()
        new_val = change.get("newValue", "").strip()
        unit    = change.get("unit", "").strip()

        if not old_val or not new_val:
            continue

        # Build search terms in priority order: longest compound first, then shorter, then bare.
        # This prevents matching "155" and leaving a unit suffix (e.g. "g") behind.
        search_terms = []
        if unit:
            search_terms.append(f"{old_val}{unit}")    # "155gm"
            search_terms.append(f"{old_val} {unit}")   # "155 gm"
            # Also try all unit prefixes (e.g. "g" from "gm", "kg" from "kgs")
            for prefix_len in range(len(unit) - 1, 0, -1):
                prefix = unit[:prefix_len]
                search_terms.append(f"{old_val}{prefix}")    # "155g"
                search_terms.append(f"{old_val} {prefix}")   # "155 g"
        search_terms.append(old_val)                    # "155" (bare fallback)

        for term in search_terms:
            pattern = re.compile(re.escape(term), re.IGNORECASE)
            if not pattern.search(full_text):
                continue

            if mode == "annotate_original":
                # Highlight old value in YELLOW, keep text unchanged
                _apply_highlight_in_paragraph(para, term, YELLOW_HIGHLIGHT, None, keep_text=True)
            elif mode == "modify_green":
                # Replace old value with new value, highlight NEW value in GREEN
                if unit and not new_val.lower().endswith(unit.lower()):
                    replacement = f"{new_val} {unit}"
                else:
                    replacement = new_val
                _apply_highlight_in_paragraph(para, term, GREEN_HIGHLIGHT, replacement, keep_text=False)
            elif mode == "modify_clean":
                # Replace old value with new value, NO highlight
                if unit and not new_val.lower().endswith(unit.lower()):
                    replacement = f"{new_val} {unit}"
                else:
                    replacement = new_val
                _apply_highlight_in_paragraph(para, term, None, replacement, keep_text=False)

            changed = True
            full_text = get_paragraph_full_text(para)
            break  # Move to next change after applying this one

    return changed


def process_table(table, changes, mode):
    """Process all cells in a table."""
    changed = False
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if process_paragraph(para, changes, mode):
                    changed = True
    return changed


def process_document(doc, changes, mode):
    """Process all paragraphs and tables in the document."""
    total_changes = 0

    for para in doc.paragraphs:
        if process_paragraph(para, changes, mode):
            total_changes += 1

    for table in doc.tables:
        if process_table(table, changes, mode):
            total_changes += 1

    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    if process_paragraph(para, changes, mode):
                        total_changes += 1
                for table in hdr.tables:
                    if process_table(table, changes, mode):
                        total_changes += 1

    return total_changes


def main():
    if len(sys.argv) < 5:
        print("Usage: wordModifier.py <mode> <input.docx> <output.docx> <changes_json>", file=sys.stderr)
        sys.exit(1)

    mode = sys.argv[1]
    input_path = sys.argv[2]
    output_path = sys.argv[3]
    changes_json = sys.argv[4]

    if mode not in ("annotate_original", "modify_green", "modify_clean"):
        print(f"Unknown mode: {mode}", file=sys.stderr)
        sys.exit(1)

    try:
        changes = json.loads(changes_json)
    except json.JSONDecodeError as e:
        print(f"Invalid changes JSON: {e}", file=sys.stderr)
        sys.exit(1)

    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Failed to open document: {e}", file=sys.stderr)
        sys.exit(1)

    total = process_document(doc, changes, mode)

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Failed to save document: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"OK:{total}", flush=True)
    sys.exit(0 if total > 0 else 2)


if __name__ == "__main__":
    main()
